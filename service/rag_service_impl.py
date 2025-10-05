import os
import uuid
from datetime import datetime
from typing import List, Dict, Any, Optional
import tempfile
from pathlib import Path

from entity.metadata_document import MetadataDocument
from entity.fragment_document import FragmentDocument
from repository.metadata_document_repository import MetadataDocumentRepository
from repository.fragment_document_repository import FragmentDocumentRepository


class RAGServiceImpl:
    """
    Implementation of RAG service for document processing and retrieval.
    Handles document upload, processing, and vector search.
    """
    
    def __init__(self):
        self.metadata_repository = MetadataDocumentRepository()
        self.fragment_repository = FragmentDocumentRepository()
        
    def process_documents(self, files, document_type: str, specialty: str, description: str = "") -> Dict[str, Any]:
        """
        Process uploaded documents and store them in the RAG system
        """
        try:
            results = {
                "processed_documents": [],
                "errors": [],
                "total_fragments": 0
            }
            
            for file in files:
                if file.filename == '':
                    continue
                    
                try:
                    # Process single document
                    doc_result = self._process_single_document(
                        file, document_type, specialty, description
                    )
                    results["processed_documents"].append(doc_result)
                    results["total_fragments"] += doc_result["fragment_count"]
                    
                except Exception as e:
                    results["errors"].append({
                        "filename": file.filename,
                        "error": str(e)
                    })
            
            return results
            
        except Exception as e:
            return {
                "processed_documents": [],
                "errors": [{"general": str(e)}],
                "total_fragments": 0
            }
    
    def _process_single_document(self, file, document_type: str, specialty: str, description: str) -> Dict[str, Any]:
        """Process a single document file"""
        
        # Generate unique document ID
        document_id = f"doc_{uuid.uuid4().hex[:12]}"
        
        # Get file info
        filename = file.filename
        file_extension = Path(filename).suffix.lower()
        
        # Save file temporarily for processing
        with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension) as temp_file:
            file.save(temp_file.name)
            temp_path = temp_file.name
        
        try:
            # Extract text from document
            text_content = self._extract_text_from_file(temp_path, file_extension)
            
            # Create chunks/fragments
            fragments = self._create_text_fragments(text_content, document_id)
            
            # Create metadata document
            metadata_doc = MetadataDocument(
                document_id=document_id,
                title=filename,
                description=description,
                document_type=document_type,
                specialty=specialty,
                file_path=temp_path,  # In production, store in proper file storage
                file_extension=file_extension,
                fragment_count=len(fragments),
                created_date=datetime.now()
            )
            
            # Save metadata
            metadata_id = self.metadata_repository.save(metadata_doc.to_dict())
            
            # Save fragments
            fragment_ids = []
            for fragment in fragments:
                fragment_id = self.fragment_repository.save(fragment.to_dict())
                fragment_ids.append(fragment_id)
            
            return {
                "document_id": document_id,
                "metadata_id": metadata_id,
                "fragment_ids": fragment_ids,
                "fragment_count": len(fragments),
                "filename": filename,
                "file_extension": file_extension
            }
            
        finally:
            # Clean up temporary file
            try:
                os.unlink(temp_path)
            except:
                pass
    
    def _extract_text_from_file(self, file_path: str, file_extension: str) -> str:
        """Extract text content from various file formats"""
        
        try:
            if file_extension == '.txt':
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            
            elif file_extension == '.pdf':
                try:
                    import PyPDF2
                    with open(file_path, 'rb') as f:
                        reader = PyPDF2.PdfReader(f)
                        text = ""
                        for page in reader.pages:
                            text += page.extract_text() + "\n"
                        return text
                except ImportError:
                    # Fallback using pdfplumber if available
                    try:
                        import pdfplumber
                        with pdfplumber.open(file_path) as pdf:
                            text = ""
                            for page in pdf.pages:
                                page_text = page.extract_text()
                                if page_text:
                                    text += page_text + "\n"
                            return text
                    except ImportError:
                        raise Exception("PDF processing libraries not available. Install PyPDF2 or pdfplumber.")
            
            elif file_extension in ['.doc', '.docx']:
                try:
                    import docx
                    doc = docx.Document(file_path)
                    text = ""
                    for paragraph in doc.paragraphs:
                        text += paragraph.text + "\n"
                    return text
                except ImportError:
                    raise Exception("python-docx library not available for DOC/DOCX files.")
            
            elif file_extension == '.md':
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            
            else:
                raise Exception(f"Unsupported file format: {file_extension}")
                
        except Exception as e:
            raise Exception(f"Error extracting text from {file_extension}: {str(e)}")
    
    def _create_text_fragments(self, text: str, document_id: str) -> List[FragmentDocument]:
        """Create text fragments from document content"""
        
        # Simple chunking strategy - split by paragraphs and limit size
        paragraphs = text.split('\n\n')
        fragments = []
        
        current_chunk = ""
        chunk_size = 1000  # characters
        overlap = 200  # character overlap
        
        for paragraph in paragraphs:
            paragraph = paragraph.strip()
            if not paragraph:
                continue
            
            # If adding this paragraph exceeds chunk size, save current chunk
            if len(current_chunk) + len(paragraph) > chunk_size and current_chunk:
                fragment = self._create_fragment(current_chunk, document_id, len(fragments))
                fragments.append(fragment)
                
                # Start new chunk with overlap
                words = current_chunk.split()
                overlap_words = words[-20:] if len(words) > 20 else words  # Last 20 words for overlap
                current_chunk = " ".join(overlap_words) + " " + paragraph
            else:
                current_chunk += " " + paragraph if current_chunk else paragraph
        
        # Add final chunk
        if current_chunk:
            fragment = self._create_fragment(current_chunk, document_id, len(fragments))
            fragments.append(fragment)
        
        return fragments
    
    def _create_fragment(self, content: str, document_id: str, sequence_number: int) -> FragmentDocument:
        """Create a single fragment document"""
        
        fragment_id = f"frag_{uuid.uuid4().hex[:12]}"
        
        # Generate embeddings (placeholder - will implement with actual embeddings)
        embeddings = self._generate_embeddings(content)
        
        return FragmentDocument(
            fragment_id=fragment_id,
            document_id=document_id,
            content=content,
            sequence_number=sequence_number,
            embeddings=embeddings,
            created_date=datetime.now()
        )
    
    def _generate_embeddings(self, text: str) -> List[float]:
        """
        Generate embeddings for text content
        TODO: Implement with actual embedding model (OpenAI, sentence-transformers, etc.)
        """
        try:
            # Try OpenAI embeddings first
            import openai
            openai.api_key = os.getenv('OPENAI_API_KEY') or os.getenv('OPENAI_KEY')
            
            if openai.api_key:
                response = openai.embeddings.create(
                    model="text-embedding-3-small",
                    input=text[:8000]  # Limit text length
                )
                return response.data[0].embedding
            else:
                # Fallback: return placeholder embeddings
                return [0.0] * 1536  # OpenAI embedding dimension
                
        except Exception as e:
            print(f"Error generating embeddings: {e}")
            # Return placeholder embeddings
            return [0.0] * 1536
    
    def search_similar_documents(self, query: str, limit: int = 5) -> List[Dict[str, Any]]:
        """
        Search for similar documents using vector similarity
        """
        try:
            # Generate query embeddings
            query_embeddings = self._generate_embeddings(query)
            
            # Search fragments (simplified - in production use proper vector search)
            all_fragments = self.fragment_repository.find_all()
            
            # Calculate similarity and rank (simplified)
            scored_fragments = []
            for fragment in all_fragments:
                # Simple similarity calculation (cosine similarity would be better)
                similarity_score = self._calculate_similarity(query, fragment.get('content', ''))
                
                scored_fragments.append({
                    'fragment': fragment,
                    'score': similarity_score
                })
            
            # Sort by score and limit results
            scored_fragments.sort(key=lambda x: x['score'], reverse=True)
            top_fragments = scored_fragments[:limit]
            
            # Get metadata for each fragment
            results = []
            for item in top_fragments:
                fragment = item['fragment']
                document_id = fragment.get('document_id')
                
                # Get document metadata
                metadata_docs = self.metadata_repository.find_by_document_id(document_id)
                metadata = metadata_docs[0] if metadata_docs else {}
                
                results.append({
                    'content': fragment.get('content', ''),
                    'score': item['score'],
                    'document_title': metadata.get('title', 'Unknown'),
                    'document_type': metadata.get('document_type', 'unknown'),
                    'specialty': metadata.get('specialty', 'general')
                })
            
            return results
            
        except Exception as e:
            print(f"Error searching documents: {e}")
            return []
    
    def _calculate_similarity(self, query: str, content: str) -> float:
        """
        Simple text similarity calculation
        TODO: Replace with proper vector similarity
        """
        query_words = set(query.lower().split())
        content_words = set(content.lower().split())
        
        if not query_words or not content_words:
            return 0.0
        
        intersection = query_words.intersection(content_words)
        union = query_words.union(content_words)
        
        return len(intersection) / len(union) if union else 0.0
    
    def get_all_documents(self) -> List[Dict[str, Any]]:
        """Get all stored documents with metadata"""
        try:
            return self.metadata_repository.find_all()
        except Exception as e:
            print(f"Error getting documents: {e}")
            return []
    
    def delete_document(self, document_id: str) -> bool:
        """Delete a document and all its fragments"""
        try:
            # Delete all fragments for this document
            self.fragment_repository.delete_by_document_id(document_id)
            
            # Delete metadata document
            metadata_docs = self.metadata_repository.find_by_document_id(document_id)
            for doc in metadata_docs:
                if doc.get('id'):
                    self.metadata_repository.delete(doc['id'])
            
            return True
            
        except Exception as e:
            print(f"Error deleting document: {e}")
            return False